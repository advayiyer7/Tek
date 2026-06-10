"""Per-format text extraction. Returns plain text or None if unreadable."""

from __future__ import annotations

import logging
from pathlib import Path

log = logging.getLogger(__name__)

MAX_CHARS = 2_000_000  # hard cap per document to bound memory


def extract_text(path: str) -> str | None:
    p = Path(path)
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(p)
        if ext == ".docx":
            return _extract_docx(p)
        return _extract_plain(p)
    except Exception as exc:  # noqa: BLE001 — a bad file must never kill indexing
        log.warning("extract failed for %s: %s", path, exc)
        return None


def _extract_plain(p: Path) -> str | None:
    raw = p.read_bytes()[: MAX_CHARS * 4]
    if b"\x00" in raw[:8192]:
        return None  # binary masquerading as text
    for encoding in ("utf-8-sig", "utf-16", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)[:MAX_CHARS]
        except (UnicodeDecodeError, UnicodeError):
            continue
    return None


def _extract_pdf(p: Path) -> str | None:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    if reader.is_encrypted:
        return None
    pages: list[str] = []
    total = 0
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
        total += len(text)
        if total > MAX_CHARS:
            break
    text = "\n\n".join(pages).strip()
    return text[:MAX_CHARS] or None


def _extract_docx(p: Path) -> str | None:
    import docx

    document = docx.Document(str(p))
    parts = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    text = "\n".join(part for part in parts if part.strip()).strip()
    return text[:MAX_CHARS] or None
